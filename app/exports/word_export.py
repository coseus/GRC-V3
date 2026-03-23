from __future__ import annotations

from collections import defaultdict
from datetime import datetime
from io import BytesIO
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def _add_logo(doc: Document, logo_path: str | None):
    if not logo_path:
        return

    p = Path(logo_path)
    if not p.exists() or not p.is_file():
        return

    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run()
        run.add_picture(str(p), width=Cm(4.5))
    except Exception:
        return


def _set_heading_style(run, size=16, color="1F4E78", bold=True):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _criticality_color(value: str):
    value = str(value or "").strip().title()
    return {
        "Critical": "C00000",
        "High": "E26B0A",
        "Medium": "BF9000",
        "Low": "38761D",
    }.get(value, "1F1F1F")


def _make_chart(domain_scores: dict):
    if not domain_scores:
        return None

    labels = list(domain_scores.keys())
    values = list(domain_scores.values())

    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.bar(labels, values)
    ax.set_title("Domain Scores")
    ax.set_ylabel("Score")
    ax.set_ylim(0, 100)
    ax.grid(axis="y", alpha=0.3)
    plt.xticks(rotation=35, ha="right")
    plt.tight_layout()

    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_word_report(
    *,
    company_name: str,
    assessment_name: str,
    framework_name: str,
    responses,
    domain_scores,
    executive_summary=None,
    recommendations=None,
    findings=None,
    logo_path: str | None = None,
    include_annex: bool = True,
    auditor_name: str | None = None,
    report_version: str | None = None,
    report_date: str | None = None,
):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    _add_logo(doc, logo_path)

    p = doc.add_paragraph()
    r = p.add_run("GRC Assessment Report")
    _set_heading_style(r, size=22)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Table Grid"
    meta_rows = [
        ("Company", company_name),
        ("Assessment", assessment_name),
        ("Framework", framework_name),
        ("Auditor", auditor_name or "-"),
        ("Report Version", report_version or "1.0"),
        ("Report Date", report_date or datetime.now().strftime("%Y-%m-%d")),
    ]
    for idx, (k, v) in enumerate(meta_rows):
        meta_table.rows[idx].cells[0].text = k
        meta_table.rows[idx].cells[1].text = v

    doc.add_paragraph(
        "This report provides an executive overview of assessment results, domain scores, findings, recommendations, and a detailed annex of evaluated controls."
    )

    doc.add_page_break()

    h = doc.add_paragraph()
    _set_heading_style(h.add_run("Domain Scores"), size=16)

    if domain_scores:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Domain"
        hdr[1].text = "Score"

        for domain_name, score in domain_scores.items():
            row = table.add_row().cells
            row[0].text = str(domain_name)
            row[1].text = f"{score:.1f}" if isinstance(score, (int, float)) else str(score)

        chart_buf = _make_chart(domain_scores)
        if chart_buf:
            doc.add_paragraph("")
            doc.add_picture(chart_buf, width=Cm(16))
    else:
        doc.add_paragraph("No domain scores available.")

    h = doc.add_paragraph()
    _set_heading_style(h.add_run("Executive Summary"), size=16)

    if executive_summary:
        for label, value in [
            ("Summary", getattr(executive_summary, "summary_text", "") or ""),
            ("Strengths", getattr(executive_summary, "strengths_text", "") or ""),
            ("Gaps", getattr(executive_summary, "gaps_text", "") or ""),
            ("Management Recommendations", getattr(executive_summary, "recommendations_text", "") or ""),
        ]:
            if value:
                p = doc.add_paragraph()
                _set_heading_style(p.add_run(label), size=12, color="1F4E78")
                doc.add_paragraph(value)
    else:
        doc.add_paragraph("No executive summary available.")

    h = doc.add_paragraph()
    _set_heading_style(h.add_run("Findings by Criticality"), size=16)

    if findings:
        grouped = defaultdict(list)
        for item in findings:
            grouped[str(item.get("criticality", "Medium")).title()].append(item)

        for criticality in ["Critical", "High", "Medium", "Low"]:
            items = grouped.get(criticality, [])
            if not items:
                continue

            p = doc.add_paragraph()
            _set_heading_style(
                p.add_run(criticality),
                size=12,
                color=_criticality_color(criticality),
            )

            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "ID"
            hdr[1].text = "Title"
            hdr[2].text = "Domain"
            hdr[3].text = "Score"
            hdr[4].text = "Notes"

            for idx, item in enumerate(items[:50], start=1):
                finding_id = f"{criticality[:1].upper()}-{idx:03d}"
                row = table.add_row().cells
                row[0].text = finding_id
                row[1].text = str(item.get("title", "") or "")
                row[2].text = str(item.get("domain", "") or "")
                row[3].text = "" if item.get("score") is None else str(item.get("score"))
                row[4].text = str(item.get("notes", "") or "")
    else:
        doc.add_paragraph("No findings included.")

    h = doc.add_paragraph()
    _set_heading_style(h.add_run("Recommendations"), size=16)

    if recommendations:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Priority"
        hdr[1].text = "Title"
        hdr[2].text = "Description"
        hdr[3].text = "Domain"

        for rec in recommendations:
            row = table.add_row().cells
            row[0].text = str(getattr(rec, "priority", "") or "")
            row[1].text = str(getattr(rec, "title", "") or "")
            row[2].text = str(getattr(rec, "description", "") or "")
            row[3].text = str(getattr(rec, "domain_name", "") or "")
    else:
        doc.add_paragraph("No recommendations available.")

    if include_annex:
        doc.add_page_break()
        h = doc.add_paragraph()
        _set_heading_style(h.add_run("Annex - Detailed Controls and Scores"), size=16)

        if responses:
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Domain"
            hdr[1].text = "Question ID"
            hdr[2].text = "Question"
            hdr[3].text = "Answer"
            hdr[4].text = "Score"
            hdr[5].text = "Risk"
            hdr[6].text = "Weight"

            for item in responses:
                row = table.add_row().cells
                row[0].text = str(item.get("domain_name") or item.get("domain") or "")
                row[1].text = str(item.get("question_id") or item.get("question_code") or "")
                row[2].text = str(item.get("question_text") or item.get("question") or "")
                row[3].text = str(item.get("answer_value") or item.get("selected_value") or "")
                row[4].text = "" if item.get("score") is None else str(item.get("score"))
                row[5].text = str(item.get("risk") or "")
                row[6].text = str(item.get("weight") or "")
        else:
            doc.add_paragraph("No annex data available.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()